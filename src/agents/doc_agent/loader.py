import os
import docx
import logging

logger = logging.getLogger(__name__)

def load_docx(file_path: str) -> str:
    """
    Loads text from an MS Word document.
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return ""
    
    try:
        doc = docx.Document(file_path)
        full_text = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                full_text.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))  # Join cell texts with a separator

        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error loading docx {file_path}: {e}")
        return ""
